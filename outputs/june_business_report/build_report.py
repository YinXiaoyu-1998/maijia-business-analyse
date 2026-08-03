from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT_DIR = Path("/Users/zhangdagang/Desktop/MAIJIA_ANALYSE/outputs/june_business_report")
DOCX_PATH = OUT_DIR / "麦家小馆_2026年6月阶段经营分析报告.docx"
MD_PATH = OUT_DIR / "麦家小馆_2026年6月阶段经营分析报告.md"

BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(95, 95, 95)


def set_run_font(run, name="Arial", size=11, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.3, color=INK, bold=True)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=9.1, color=INK)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_metric_strip(doc):
    metrics = [
        ("营业收入", "868,395.83 元"),
        ("营业额", "983,016.26 元"),
        ("订单量", "11,795 单"),
        ("折后单均", "73.62 元"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(metrics):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F4F7FB")
        set_cell_margins(cell, top=150, bottom=150, start=160, end=160)
        cell.width = Inches(1.58)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "\n")
        set_run_font(r1, size=9.5, color=MUTED, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=13, color=BLUE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_docx():
    doc = Document()
    format_doc(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("麦家小馆 | 2026年6月阶段经营分析")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("数据来源：美团管家营业概览导出；当前数据范围 2026/06/01-2026/06/14")
    set_run_font(fr, size=8.5, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("麦家小馆 2026年6月阶段经营分析报告")
    set_run_font(tr, size=22, color=RGBColor(0, 0, 0), bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sr = subtitle.add_run("分析范围：2026/06/01-2026/06/14；拆分口径：第一周 06/01-06/07，第二周 06/08-06/14")
    set_run_font(sr, size=10.5, color=MUTED)

    add_metric_strip(doc)

    doc.add_heading("一、总体判断", level=1)
    add_para(
        doc,
        "6月前两周整体营业收入 868,395.83 元，营业额 983,016.26 元，优惠金额 114,620.43 元，优惠占比 11.66%，订单量 11,795 单。折后单均 73.62 元，折后人均 55.82 元，单店日均营业收入 20,676.09 元。",
    )
    add_para(
        doc,
        "第二周比第一周明显走强：营业收入增加 33,626.21 元，环比 +8.06%；营业额增加 39,663.54 元，环比 +8.41%。订单量仅增加 81 单，环比 +1.38%，说明增长主要来自客单和人均提升，而不是单量扩张。",
    )
    add_para(
        doc,
        "需要重点关注优惠压力：第二周优惠金额增加 6,037.33 元，环比 +11.12%，快于收入增速；优惠占比从 11.51%升至 11.80%。增长质量整体较好，但折扣效率需要持续盯紧。",
    )

    doc.add_heading("二、周度核心指标", level=1)
    add_table(
        doc,
        ["指标", "第一周", "第二周", "环比变化"],
        [
            ["营业收入", "417,384.81", "451,011.02", "+33,626.21 / +8.06%"],
            ["营业额", "471,676.36", "511,339.90", "+39,663.54 / +8.41%"],
            ["优惠金额", "54,291.55", "60,328.88", "+6,037.33 / +11.12%"],
            ["优惠占比", "11.51%", "11.80%", "+0.29pct"],
            ["订单量", "5,857", "5,938", "+81 / +1.38%"],
            ["折后单均", "71.26", "75.95", "+4.69 / +6.58%"],
            ["折后人均", "53.72", "57.91", "+4.19 / +7.80%"],
            ["单店日均营业收入", "19,875.47", "21,476.72", "+1,601.25 / +8.06%"],
        ],
        widths=[1.6, 1.35, 1.35, 2.05],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("三、门店表现", level=1)
    add_para(doc, "第二周增长主要来自苏州街店和通州保利店，常营店收入略有回落。")
    add_table(
        doc,
        ["门店", "第一周收入", "第二周收入", "环比变化", "第二周订单", "第二周折后单均"],
        [
            ["苏州街店", "146,774.43", "165,655.83", "+18,881.40 / +12.86%", "1,974", "83.92"],
            ["常营店", "147,833.22", "146,414.59", "-1,418.63 / -0.96%", "2,238", "65.42"],
            ["通州保利店", "122,777.16", "138,940.60", "+16,163.44 / +13.16%", "1,726", "80.50"],
        ],
        widths=[1.15, 1.1, 1.1, 1.55, 1.0, 1.05],
        header_fill=LIGHT_BLUE,
    )
    add_bullet(doc, "苏州街店与通州保利店是第二周增长引擎，且都体现为客单提升。")
    add_bullet(doc, "常营店订单最高但单均最低，第二周优惠占比升至 13.61%，收入反而微降，应优先复盘低客单和折扣结构。")

    doc.add_heading("四、渠道结构", level=1)
    add_para(doc, "堂食仍是绝对主力。前两周店内销售营业收入 741,157.94 元，占总收入 85.35%；外卖平台合计 127,237.89 元，占 14.65%。")
    add_table(
        doc,
        ["渠道", "第一周收入", "第二周收入", "环比变化", "第二周优惠占比"],
        [
            ["店内销售", "355,077.37", "386,080.57", "+31,003.20 / +8.73%", "5.43%"],
            ["美团外卖", "44,583.81", "47,515.89", "+2,932.08 / +6.58%", "38.33%"],
            ["淘宝闪购", "13,941.66", "13,327.45", "-614.21 / -4.41%", "33.52%"],
            ["京东秒送", "3,781.97", "4,087.11", "+305.14 / +8.07%", "31.73%"],
        ],
        widths=[1.25, 1.25, 1.25, 1.55, 1.25],
        header_fill=LIGHT_BLUE,
    )
    add_para(doc, "外卖端贡献收入 14.65%，但外卖相关优惠金额 75,120.77 元，占总优惠 65.54%，是利润压力的核心来源。")

    doc.add_heading("五、品类与菜品", level=1)
    add_para(doc, "前两周部门收入前三为烤串 264,876.88 元、小吃 147,542.04 元、刀削面 108,217.93 元，三类合计占总营业收入 59.97%；加上凉菜后，前四类占 70.50%。经营结构高度集中，烤串是基本盘。畅销菜从第一周“招牌刀削面”转向第二周“高品质大串羔羊肉”，精酿提升也支撑了客单增长。")
    add_para(doc, "第二周主要增量来自烤串、凉菜、精酿啤酒类和小吃；回落项集中在刀削面与手擀面。")

    doc.add_heading("六、会员与敏感操作", level=1)
    add_para(doc, "会员营业额占比从第一周 28.98%提升到第二周 30.31%，前两周整体 29.68%。会员占比提升是稳定复购的正向信号，后续可继续围绕晚市、烤串、精酿组合做会员触达。")
    add_para(doc, "敏感操作方面，第二周退菜 398 菜，比第一周增加 28 菜，环比 +7.57%；撤单 48 单，比第一周减少 6 单，环比 -11.11%。前两周退菜合计 768 菜，撤单合计 102 单。")

    doc.add_heading("七、下阶段动作建议", level=1)
    for item in [
        "复制高客单打法：复盘苏州街店、通州保利店第二周高客单来源，拆到时段、品类、组合和员工推荐动作，再移植到常营店。",
        "常营店专项整改：优先查低客单订单、套餐结构、折扣使用和退菜高发菜品，目标是先把折后单均拉回到 70 元以上。",
        "外卖控补贴：按平台拆满减、配送补贴、服务费和菜品毛利，先压缩 30% 以上优惠率中的低效部分。",
        "放大精酿与烤串组合：精酿啤酒第二周增长突出，可作为晚市客单提升抓手，但要同步控制第二周 24.27% 的精酿优惠占比。",
        "退菜复盘到班次和菜品：常营店退菜偏高，建议按菜品、班次、员工、原因分类，先处理前 5 个高频问题。",
    ]:
        add_number(doc, item)

    doc.add_heading("八、数据文件与口径", level=1)
    add_bullet(doc, "总表：永杰厚道_营业概览_20260601-20260614.xlsx")
    add_bullet(doc, "第一周：永杰厚道_营业概览_20260601-20260607.xlsx")
    add_bullet(doc, "第二周：永杰厚道_营业概览_20260608-20260614.xlsx")
    add_bullet(doc, "说明：当前目录未发现 2026/06/15-2026/06/30 数据，本报告按 6 月阶段经营分析处理；补齐下半月数据后可升级为完整月报。")

    doc.save(DOCX_PATH)


MD = """# 麦家小馆 2026年6月阶段经营分析报告

数据范围：2026/06/01-2026/06/14  
拆分口径：第一周 2026/06/01-2026/06/07；第二周 2026/06/08-2026/06/14  
数据来源：美团管家「报表中心 - 经营分析 - 营业概览」导出  
说明：当前目录未发现 2026/06/15-2026/06/30 数据，本报告按 6 月阶段经营分析处理。

## 一、总体判断

6月前两周整体营业收入 868,395.83 元，营业额 983,016.26 元，优惠金额 114,620.43 元，优惠占比 11.66%，订单量 11,795 单。折后单均 73.62 元，折后人均 55.82 元，单店日均营业收入 20,676.09 元。

第二周比第一周明显走强：营业收入增加 33,626.21 元，环比 +8.06%；营业额增加 39,663.54 元，环比 +8.41%。订单量只增加 81 单，环比 +1.38%，说明增长主要来自客单和人均提升，而不是单量扩张。

需要注意的是，优惠金额第二周增加 6,037.33 元，环比 +11.12%，快于收入增速；优惠占比从 11.51%升至 11.80%。增长质量整体较好，但折扣压力在抬头。

## 二、周度核心指标

| 指标 | 第一周 | 第二周 | 环比变化 |
|---|---:|---:|---:|
| 营业收入 | 417,384.81 | 451,011.02 | +33,626.21 / +8.06% |
| 营业额 | 471,676.36 | 511,339.90 | +39,663.54 / +8.41% |
| 优惠金额 | 54,291.55 | 60,328.88 | +6,037.33 / +11.12% |
| 优惠占比 | 11.51% | 11.80% | +0.29pct |
| 订单量 | 5,857 | 5,938 | +81 / +1.38% |
| 折后单均 | 71.26 | 75.95 | +4.69 / +6.58% |
| 折后人均 | 53.72 | 57.91 | +4.19 / +7.80% |
| 单店日均营业收入 | 19,875.47 | 21,476.72 | +1,601.25 / +8.06% |

## 三、门店表现

第二周增长主要来自苏州街店和通州保利店，常营店收入略有回落。

| 门店 | 第一周收入 | 第二周收入 | 环比变化 | 第二周订单 | 第二周折后单均 |
|---|---:|---:|---:|---:|---:|
| 苏州街店 | 146,774.43 | 165,655.83 | +18,881.40 / +12.86% | 1,974 | 83.92 |
| 常营店 | 147,833.22 | 146,414.59 | -1,418.63 / -0.96% | 2,238 | 65.42 |
| 通州保利店 | 122,777.16 | 138,940.60 | +16,163.44 / +13.16% | 1,726 | 80.50 |

判断：苏州街店和通州保利店是第二周增长引擎，且都体现为客单提升；常营店订单最高但单均最低，第二周优惠占比升至 13.61%，收入反而微降，需要重点看套餐、折扣和低客单订单结构。

## 四、渠道结构

堂食仍是绝对主力。6月前两周店内销售营业收入 741,157.94 元，占总收入 85.35%；外卖平台合计 127,237.89 元，占 14.65%。

| 渠道 | 第一周收入 | 第二周收入 | 环比变化 | 第二周优惠占比 |
|---|---:|---:|---:|---:|
| 店内销售 | 355,077.37 | 386,080.57 | +31,003.20 / +8.73% | 5.43% |
| 美团外卖 | 44,583.81 | 47,515.89 | +2,932.08 / +6.58% | 38.33% |
| 淘宝闪购 | 13,941.66 | 13,327.45 | -614.21 / -4.41% | 33.52% |
| 京东秒送 | 3,781.97 | 4,087.11 | +305.14 / +8.07% | 31.73% |

判断：增长主要在店内销售，外卖端增长较小且优惠率长期处在 31%-38%高位。外卖贡献收入 14.65%，但外卖相关优惠金额 75,120.77 元，占总优惠 65.54%，是利润压力的核心来源。

## 五、品类与菜品

前两周部门收入前三为烤串 264,876.88 元、小吃 147,542.04 元、刀削面 108,217.93 元，三类合计占总营业收入 59.97%；加上凉菜后，前四类占 70.50%。经营结构高度集中，烤串是基本盘。

第二周品类增量较明显的方向：

- 烤串：+13,320.84 元，环比 +10.59%
- 凉菜：+5,558.09 元，环比 +12.89%
- 精酿啤酒类：+5,445.88 元，环比 +26.66%
- 小吃：+3,371.82 元，环比 +4.68%

第二周回落的方向：

- 刀削面：-1,124.33 元，环比 -2.06%
- 手擀面：-491.99 元，环比 -1.58%

畅销菜方面，第一周第一是招牌刀削面 27,560.00 元，第二周第一变为高品质大串羔羊肉 27,937.00 元；精酿黄啤、白啤第二周提升明显，和第二周客单、人均提升一致。

## 六、会员与敏感操作

会员营业额占比从第一周 28.98%提升到第二周 30.31%，前两周整体 29.68%。会员占比提升对复购和稳定客群是好信号，但还可以继续拉高。

敏感操作方面，第二周退菜 398 菜，比第一周增加 28 菜，环比 +7.57%；撤单 48 单，比第一周减少 6 单，环比 -11.11%。前两周退菜合计 768 菜，撤单合计 102 单。

## 七、下阶段动作建议

1. 复制高客单打法：复盘苏州街店、通州保利店第二周高客单来源，拆到时段、品类、组合和员工推荐动作，再移植到常营店。
2. 常营店专项整改：优先查低客单订单、套餐结构、折扣使用和退菜高发菜品，目标是先把折后单均拉回到 70 元以上。
3. 外卖控补贴：按平台拆满减、配送补贴、服务费和菜品毛利，先压缩 30% 以上优惠率中的低效部分。
4. 放大精酿与烤串组合：精酿啤酒第二周增长突出，可作为晚市客单提升抓手，但要同步控制第二周 24.27% 的精酿优惠占比。
5. 退菜复盘到班次和菜品：常营店退菜偏高，建议按菜品、班次、员工、原因分类，先处理前 5 个高频问题。

## 八、数据文件

- 总表：永杰厚道_营业概览_20260601-20260614.xlsx
- 第一周：永杰厚道_营业概览_20260601-20260607.xlsx
- 第二周：永杰厚道_营业概览_20260608-20260614.xlsx
"""


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    MD_PATH.write_text(MD, encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
